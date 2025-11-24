"""
DOCX Parser Module

Specialized DOCX document processing using python-docx.

Features:
- Text extraction with formatting preservation
- Heading and structure extraction
- Image extraction
- Table extraction with cell data
- Document metadata extraction
- Style and formatting information

Dependencies:
- python-docx >= 1.1.0
"""

from typing import Dict, List, Optional, Any
import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from app.documents.config import DocumentConfig
from app.logger import get_logger

logger = get_logger(__name__)


class DOCXParserError(Exception):
    """Base exception for DOCX parsing errors"""
    pass


class DOCXCorruptedError(DOCXParserError):
    """Raised when DOCX file is corrupted"""
    pass


class DOCXParser:
    """
    DOCX document parser using python-docx

    Handles text extraction, image extraction, table extraction,
    and metadata extraction from Microsoft Word documents.
    """

    def __init__(self, config: Optional[DocumentConfig] = None):
        """
        Initialize DOCX parser

        Args:
            config: Document processing configuration

        Raises:
            ImportError: If python-docx is not installed
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx>=1.1.0"
            )

        self.config = config or DocumentConfig()
        self.settings = self.config.docx_settings
        self.logger = logger

    async def extract_text(self, file_path: str) -> str:
        """
        Extract all text from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content

        Raises:
            DOCXCorruptedError: If DOCX is corrupted
            FileNotFoundError: If file doesn't exist
        """
        self.logger.info(f"Extracting text from DOCX: {file_path}")

        try:
            doc = self._open_docx(file_path)

            # Extract text from all paragraphs
            text_parts = []

            for element in doc.element.body:
                if isinstance(element, CT_P):
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                elif isinstance(element, CT_Tbl):
                    # Extract text from tables
                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(table_text)

            full_text = "\n\n".join(text_parts)

            # Post-process text
            if self.config.normalize_unicode:
                full_text = self._normalize_unicode(full_text)

            if not self.config.preserve_whitespace:
                full_text = self._normalize_whitespace(full_text)

            self.logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    async def extract_paragraphs(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract paragraphs with formatting information

        Args:
            file_path: Path to DOCX file

        Returns:
            List of paragraph dicts with text and formatting
        """
        self.logger.info(f"Extracting paragraphs from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        paragraphs_data = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            para_data = {
                "text": para.text,
                "style": para.style.name if para.style else None,
            }

            if self.settings.get("preserve_formatting", True):
                para_data["formatting"] = {
                    "alignment": str(para.alignment) if para.alignment else None,
                    "is_bold": any(run.bold for run in para.runs),
                    "is_italic": any(run.italic for run in para.runs),
                    "is_underline": any(run.underline for run in para.runs),
                }

            paragraphs_data.append(para_data)

        self.logger.info(f"Extracted {len(paragraphs_data)} paragraphs")
        return paragraphs_data

    async def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract document headings with hierarchy

        Args:
            file_path: Path to DOCX file

        Returns:
            List of heading dicts with level and text
        """
        self.logger.info(f"Extracting headings from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        headings = []

        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                try:
                    # Extract heading level from style name (e.g., "Heading 1" -> 1)
                    level = int(para.style.name.split()[-1])

                    headings.append({
                        "level": level,
                        "text": para.text,
                        "style": para.style.name,
                    })
                except (ValueError, IndexError):
                    # Not a standard heading style
                    continue

        self.logger.info(f"Extracted {len(headings)} headings")
        return headings

    async def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract embedded images from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            List of image data dicts with image bytes and metadata
        """
        if not self.settings.get("extract_images", True):
            return []

        self.logger.info(f"Extracting images from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        images = []

        # Extract images from relationships
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_data = {
                        "rel_id": rel_id,
                        "bytes": rel.target_part.blob,
                        "content_type": rel.target_part.content_type,
                        "extension": self._get_image_extension(
                            rel.target_part.content_type
                        ),
                    }

                    images.append(image_data)

                except Exception as e:
                    self.logger.warning(
                        f"Failed to extract image {rel_id}: {str(e)}"
                    )

        self.logger.info(f"Extracted {len(images)} images")
        return images

    async def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from DOCX with cell data

        Args:
            file_path: Path to DOCX file

        Returns:
            List of table data with rows and cells
        """
        if not self.settings.get("extract_tables", True):
            return []

        self.logger.info(f"Extracting tables from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        tables_data = []

        for table_idx, table in enumerate(doc.tables):
            rows_data = []

            for row in table.rows:
                cells_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells_data)

            table_data = {
                "table_index": table_idx,
                "rows": rows_data,
                "num_rows": len(rows_data),
                "num_cols": len(rows_data[0]) if rows_data else 0,
            }

            tables_data.append(table_data)

        self.logger.info(f"Extracted {len(tables_data)} tables")
        return tables_data

    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract DOCX metadata

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with metadata fields
        """
        self.logger.info(f"Extracting metadata from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        # Extract core properties
        core_props = doc.core_properties

        metadata = {
            "format": "docx",
            "file_size": os.path.getsize(file_path),
            "file_name": os.path.basename(file_path),
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "comments": core_props.comments or "",
            "category": core_props.category or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision,
        }

        # Count elements
        metadata["num_paragraphs"] = len(doc.paragraphs)
        metadata["num_tables"] = len(doc.tables)
        metadata["num_sections"] = len(doc.sections)

        return metadata

    async def extract_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Extract document structure (headings hierarchy, sections, etc.)

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with structure information
        """
        self.logger.info(f"Extracting structure from DOCX: {file_path}")

        doc = self._open_docx(file_path)

        structure = {
            "headings": [],
            "sections": [],
            "styles_used": set(),
        }

        # Extract headings hierarchy
        headings = await self.extract_headings(file_path)
        structure["headings"] = headings

        # Extract section information
        for section_idx, section in enumerate(doc.sections):
            section_data = {
                "section_index": section_idx,
                "page_width": section.page_width.inches if section.page_width else None,
                "page_height": section.page_height.inches if section.page_height else None,
                "orientation": str(section.orientation),
            }
            structure["sections"].append(section_data)

        # Collect used styles
        for para in doc.paragraphs:
            if para.style:
                structure["styles_used"].add(para.style.name)

        structure["styles_used"] = sorted(list(structure["styles_used"]))

        return structure

    async def extract_comments(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract comments from DOCX (if enabled)

        Args:
            file_path: Path to DOCX file

        Returns:
            List of comment dicts

        Note:
            This is a basic implementation. Full comment extraction
            requires deeper XML parsing.
        """
        if not self.settings.get("extract_comments", False):
            return []

        self.logger.info(f"Extracting comments from DOCX: {file_path}")

        # TODO: Implement comment extraction using python-docx XML API
        # This requires parsing the comments.xml part

        return []

    def _open_docx(self, file_path: str) -> Document:
        """
        Open DOCX document with error handling

        Args:
            file_path: Path to DOCX file

        Returns:
            Opened Document

        Raises:
            FileNotFoundError: If file doesn't exist
            DOCXCorruptedError: If DOCX is corrupted
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            return Document(file_path)
        except Exception as e:
            raise DOCXCorruptedError(f"Failed to open DOCX: {str(e)}")

    def _extract_table_text(self, table: Table) -> str:
        """
        Extract text from table in a readable format

        Args:
            table: python-docx Table object

        Returns:
            Formatted table text
        """
        rows_text = []

        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells_text))

        return "\n".join(rows_text)

    def _get_image_extension(self, content_type: str) -> str:
        """
        Get file extension from content type

        Args:
            content_type: MIME type (e.g., "image/png")

        Returns:
            File extension (e.g., "png")
        """
        extensions = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
        }

        return extensions.get(content_type, "bin")

    def _normalize_unicode(self, text: str) -> str:
        """Normalize Unicode characters"""
        import unicodedata
        return unicodedata.normalize('NFKC', text)

    def _normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace in text"""
        import re
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n\n+', '\n\n', text)
        return text.strip()

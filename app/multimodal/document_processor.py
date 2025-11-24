"""
Document Processor - Text Extraction from Documents

SPECIFICATION:
- Extract text from PDF, DOCX, TXT, MD files
- Chunk documents for RAG indexing
- Extract metadata (author, title, pages, etc.)
- Handle various encodings
- Preserve document structure

ARCHITECTURE:
- PDF: PyPDF2 for text extraction
- DOCX: python-docx for Word documents
- TXT/MD: Direct text reading with encoding detection
- Chunking: Semantic chunking with overlap

PATTERN: Strategy pattern for different document types
WHY: Each document format requires specific handling
RESILIENCE: Encoding fallbacks, error handling per format
"""

import asyncio
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime
import chardet

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

from app.logger import api_logger


class DocumentProcessor:
    """
    Document text extraction and processing

    PATTERN: Service class with format-specific extractors
    WHY: Unified interface for multiple document formats
    """

    # Chunk parameters for RAG
    DEFAULT_CHUNK_SIZE = 500  # characters
    DEFAULT_CHUNK_OVERLAP = 50  # characters

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
    ):
        """
        Initialize document processor

        Args:
            chunk_size: Size of text chunks for RAG
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        api_logger.info(
            "document_processor_initialized",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            pdf_support=PdfReader is not None,
            docx_support=Document is not None
        )

    async def process_document(
        self,
        file_path: str,
        extract_metadata: bool = True,
        chunk_text: bool = True
    ) -> Dict:
        """
        Process document and extract text

        ALGORITHM:
        1. Detect document format
        2. Extract text using format-specific extractor
        3. Extract metadata if requested
        4. Chunk text if requested
        5. Return structured results

        Args:
            file_path: Path to document file
            extract_metadata: Extract document metadata
            chunk_text: Chunk text for RAG

        Returns:
            Processing results dictionary
        """
        start_time = datetime.utcnow()

        try:
            path = Path(file_path)
            extension = path.suffix.lower()

            api_logger.info(
                "document_processing_started",
                file_path=file_path,
                extension=extension
            )

            # Route to format-specific extractor
            if extension == '.pdf':
                result = await self._extract_from_pdf(file_path, extract_metadata)
            elif extension == '.docx':
                result = await self._extract_from_docx(file_path, extract_metadata)
            elif extension in ['.txt', '.md']:
                result = await self._extract_from_text(file_path, extract_metadata)
            else:
                raise ValueError(f"Unsupported document format: {extension}")

            # Chunk text if requested
            if chunk_text and result.get('success'):
                result['chunks'] = self._chunk_text(result['text'])

            # Add processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            result['processing_time_ms'] = round(processing_time, 2)

            api_logger.info(
                "document_processing_complete",
                file_path=file_path,
                success=result.get('success'),
                text_length=len(result.get('text', '')),
                chunk_count=len(result.get('chunks', [])),
                processing_time_ms=result['processing_time_ms']
            )

            return result

        except Exception as e:
            api_logger.error(
                "document_processing_error",
                file_path=file_path,
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True
            )
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }

    async def _extract_from_pdf(
        self,
        file_path: str,
        extract_metadata: bool
    ) -> Dict:
        """Extract text from PDF"""
        if PdfReader is None:
            return {
                "success": False,
                "error": "PyPDF2 not installed. Install with: pip install PyPDF2"
            }

        try:
            reader = PdfReader(file_path)

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())

            text = "\n\n".join(text_parts)

            result = {
                "success": True,
                "text": text,
                "format": "pdf",
                "page_count": len(reader.pages)
            }

            # Extract metadata
            if extract_metadata and reader.metadata:
                result['metadata'] = {
                    "title": reader.metadata.get('/Title', ''),
                    "author": reader.metadata.get('/Author', ''),
                    "subject": reader.metadata.get('/Subject', ''),
                    "creator": reader.metadata.get('/Creator', ''),
                }

            return result

        except Exception as e:
            api_logger.error("pdf_extraction_error", file_path=file_path, error=str(e))
            return {
                "success": False,
                "error": f"PDF extraction failed: {str(e)}"
            }

    async def _extract_from_docx(
        self,
        file_path: str,
        extract_metadata: bool
    ) -> Dict:
        """Extract text from DOCX"""
        if Document is None:
            return {
                "success": False,
                "error": "python-docx not installed. Install with: pip install python-docx"
            }

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text = "\n\n".join(text_parts)

            result = {
                "success": True,
                "text": text,
                "format": "docx",
                "paragraph_count": len(doc.paragraphs)
            }

            # Extract metadata
            if extract_metadata:
                core_properties = doc.core_properties
                result['metadata'] = {
                    "title": core_properties.title or '',
                    "author": core_properties.author or '',
                    "subject": core_properties.subject or '',
                    "created": core_properties.created.isoformat() if core_properties.created else None,
                    "modified": core_properties.modified.isoformat() if core_properties.modified else None,
                }

            return result

        except Exception as e:
            api_logger.error("docx_extraction_error", file_path=file_path, error=str(e))
            return {
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}"
            }

    async def _extract_from_text(
        self,
        file_path: str,
        extract_metadata: bool
    ) -> Dict:
        """Extract text from TXT/MD files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()

            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] or 'utf-8'

            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()

            path = Path(file_path)
            result = {
                "success": True,
                "text": text,
                "format": path.suffix.lower().replace('.', ''),
                "encoding": encoding,
                "line_count": len(text.split('\n'))
            }

            # Basic metadata
            if extract_metadata:
                result['metadata'] = {
                    "filename": path.name,
                    "size_bytes": len(raw_data),
                }

            return result

        except Exception as e:
            api_logger.error("text_extraction_error", file_path=file_path, error=str(e))
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}"
            }

    def _chunk_text(self, text: str) -> List[Dict]:
        """
        Chunk text for RAG indexing

        ALGORITHM: Sliding window with overlap
        WHY: Maintain context between chunks

        Args:
            text: Full document text

        Returns:
            List of chunk dictionaries
        """
        if not text:
            return []

        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            # Extract chunk
            end = start + self.chunk_size
            chunk_text = text[start:end]

            # Try to end at sentence boundary
            if end < len(text):
                # Look for sentence ending
                for delimiter in ['. ', '.\n', '! ', '?\n']:
                    last_delimiter = chunk_text.rfind(delimiter)
                    if last_delimiter > self.chunk_size * 0.7:  # At least 70% of chunk
                        chunk_text = chunk_text[:last_delimiter + 1]
                        end = start + last_delimiter + 1
                        break

            chunks.append({
                "chunk_id": chunk_id,
                "text": chunk_text.strip(),
                "start_char": start,
                "end_char": end,
                "length": len(chunk_text.strip())
            })

            # Move to next chunk with overlap
            start = end - self.chunk_overlap
            chunk_id += 1

        api_logger.debug(
            "text_chunked",
            total_length=len(text),
            chunk_count=len(chunks),
            avg_chunk_size=sum(c['length'] for c in chunks) / len(chunks) if chunks else 0
        )

        return chunks

    async def extract_preview(self, file_path: str, max_length: int = 200) -> str:
        """
        Extract preview text from document

        Args:
            file_path: Path to document
            max_length: Maximum preview length

        Returns:
            Preview text
        """
        result = await self.process_document(file_path, extract_metadata=False, chunk_text=False)

        if result.get('success'):
            text = result.get('text', '')
            preview = text[:max_length]
            if len(text) > max_length:
                preview += "..."
            return preview

        return "[Preview not available]"


# Singleton instance
document_processor = DocumentProcessor()

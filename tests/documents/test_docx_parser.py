"""
Tests for DOCXParser

Tests DOCX document processing including:
- Text extraction
- Paragraph extraction with formatting
- Heading extraction
- Image extraction
- Table extraction
- Metadata extraction
- Structure extraction
"""

import pytest
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from app.documents import DOCXParser, DocumentConfig
from app.documents.docx_parser import DOCXParserError, DOCXCorruptedError


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
class TestDOCXParser:
    """Test DOCXParser class"""

    @pytest.fixture
    def parser(self):
        """Create parser instance"""
        config = DocumentConfig()
        return DOCXParser(config)

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create sample DOCX for testing"""
        if not PYTHON_DOCX_AVAILABLE:
            pytest.skip("python-docx not available")

        file_path = tmp_path / "sample.docx"

        # Create simple DOCX
        doc = Document()
        doc.add_heading('Main Title', level=1)
        doc.add_paragraph('This is a test DOCX document.')
        doc.add_heading('Section 1', level=2)
        doc.add_paragraph('Content in section 1.')

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'

        doc.save(str(file_path))

        return str(file_path)

    def test_init(self, parser):
        """Test parser initialization"""
        assert parser is not None
        assert parser.config is not None
        assert parser.settings is not None

    @pytest.mark.asyncio
    async def test_extract_text(self, parser, sample_docx):
        """Test text extraction"""
        text = await parser.extract_text(sample_docx)

        assert isinstance(text, str)
        assert len(text) > 0
        assert "test DOCX document" in text
        assert "Main Title" in text

    @pytest.mark.asyncio
    async def test_extract_text_nonexistent_file(self, parser):
        """Test extracting from nonexistent file"""
        with pytest.raises(FileNotFoundError):
            await parser.extract_text("/nonexistent/file.docx")

    @pytest.mark.asyncio
    async def test_extract_paragraphs(self, parser, sample_docx):
        """Test paragraph extraction"""
        paragraphs = await parser.extract_paragraphs(sample_docx)

        assert isinstance(paragraphs, list)
        assert len(paragraphs) > 0
        assert "text" in paragraphs[0]
        assert "style" in paragraphs[0]

    @pytest.mark.asyncio
    async def test_extract_paragraphs_with_formatting(self, tmp_path):
        """Test paragraph extraction with formatting"""
        if not PYTHON_DOCX_AVAILABLE:
            pytest.skip("python-docx not available")

        file_path = tmp_path / "formatted.docx"

        # Create DOCX with formatting
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        doc.save(str(file_path))

        config = DocumentConfig()
        config.docx_settings["preserve_formatting"] = True
        parser = DOCXParser(config)

        paragraphs = await parser.extract_paragraphs(str(file_path))

        assert len(paragraphs) > 0
        assert "formatting" in paragraphs[0]

    @pytest.mark.asyncio
    async def test_extract_headings(self, parser, sample_docx):
        """Test heading extraction"""
        headings = await parser.extract_headings(sample_docx)

        assert isinstance(headings, list)
        assert len(headings) >= 2
        assert "level" in headings[0]
        assert "text" in headings[0]
        assert headings[0]["text"] == "Main Title"

    @pytest.mark.asyncio
    async def test_extract_tables(self, parser, sample_docx):
        """Test table extraction"""
        tables = await parser.extract_tables(sample_docx)

        assert isinstance(tables, list)
        assert len(tables) >= 1
        assert "rows" in tables[0]
        assert "num_rows" in tables[0]
        assert "num_cols" in tables[0]
        assert tables[0]["num_rows"] == 2
        assert tables[0]["num_cols"] == 2

    @pytest.mark.asyncio
    async def test_extract_tables_disabled(self, sample_docx):
        """Test table extraction when disabled"""
        config = DocumentConfig()
        config.docx_settings["extract_tables"] = False
        parser = DOCXParser(config)

        tables = await parser.extract_tables(sample_docx)

        assert isinstance(tables, list)
        assert len(tables) == 0

    @pytest.mark.asyncio
    async def test_extract_images_disabled(self, sample_docx):
        """Test image extraction when disabled"""
        config = DocumentConfig()
        config.docx_settings["extract_images"] = False
        parser = DOCXParser(config)

        images = await parser.extract_images(sample_docx)

        assert isinstance(images, list)
        assert len(images) == 0

    @pytest.mark.asyncio
    async def test_extract_metadata(self, parser, sample_docx):
        """Test metadata extraction"""
        metadata = await parser.extract_metadata(sample_docx)

        assert isinstance(metadata, dict)
        assert metadata["format"] == "docx"
        assert "file_size" in metadata
        assert "num_paragraphs" in metadata
        assert "num_tables" in metadata
        assert metadata["num_tables"] >= 1

    @pytest.mark.asyncio
    async def test_extract_structure(self, parser, sample_docx):
        """Test structure extraction"""
        structure = await parser.extract_structure(sample_docx)

        assert isinstance(structure, dict)
        assert "headings" in structure
        assert "sections" in structure
        assert "styles_used" in structure
        assert len(structure["headings"]) >= 2

    @pytest.mark.asyncio
    async def test_extract_comments(self, parser, sample_docx):
        """Test comment extraction (disabled by default)"""
        comments = await parser.extract_comments(sample_docx)

        # Comments extraction not fully implemented
        assert isinstance(comments, list)

    def test_get_image_extension(self, parser):
        """Test image extension detection"""
        assert parser._get_image_extension("image/png") == "png"
        assert parser._get_image_extension("image/jpeg") == "jpg"
        assert parser._get_image_extension("image/gif") == "gif"
        assert parser._get_image_extension("unknown") == "bin"

    def test_normalize_unicode(self, parser):
        """Test Unicode normalization"""
        text = "café résumé"
        result = parser._normalize_unicode(text)

        assert isinstance(result, str)
        assert len(result) > 0

    def test_normalize_whitespace(self, parser):
        """Test whitespace normalization"""
        text = "Multiple    spaces\n\n\n\nand newlines"
        result = parser._normalize_whitespace(text)

        assert "    " not in result
        assert "\n\n\n" not in result


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not installed")
class TestDOCXParserAdvanced:
    """Advanced DOCX parser tests"""

    @pytest.fixture
    def complex_docx(self, tmp_path):
        """Create complex DOCX with multiple features"""
        if not PYTHON_DOCX_AVAILABLE:
            pytest.skip("python-docx not available")

        file_path = tmp_path / "complex.docx"

        doc = Document()

        # Add various heading levels
        doc.add_heading('Title', level=1)
        doc.add_heading('Section 1', level=2)
        doc.add_heading('Subsection 1.1', level=3)
        doc.add_paragraph('Content here.')

        # Add multiple tables
        for i in range(2):
            table = doc.add_table(rows=3, cols=3)
            for row in range(3):
                for col in range(3):
                    table.cell(row, col).text = f'R{row}C{col}'

        # Add more content
        doc.add_heading('Section 2', level=2)
        doc.add_paragraph('More content.')

        doc.save(str(file_path))
        return str(file_path)

    @pytest.mark.asyncio
    async def test_complex_document_processing(self, complex_docx):
        """Test processing complex document"""
        parser = DOCXParser()

        text = await parser.extract_text(complex_docx)
        headings = await parser.extract_headings(complex_docx)
        tables = await parser.extract_tables(complex_docx)
        metadata = await parser.extract_metadata(complex_docx)

        # Verify all extractions succeeded
        assert len(text) > 0
        assert len(headings) >= 4  # Title + 2 sections + 1 subsection
        assert len(tables) >= 2
        assert metadata["num_tables"] >= 2

    @pytest.mark.asyncio
    async def test_document_with_core_properties(self, tmp_path):
        """Test document with core properties set"""
        if not PYTHON_DOCX_AVAILABLE:
            pytest.skip("python-docx not available")

        file_path = tmp_path / "with_props.docx"

        doc = Document()
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.core_properties.subject = "Test Subject"
        doc.add_paragraph("Content")

        doc.save(str(file_path))

        parser = DOCXParser()
        metadata = await parser.extract_metadata(str(file_path))

        assert metadata["title"] == "Test Title"
        assert metadata["author"] == "Test Author"
        assert metadata["subject"] == "Test Subject"
